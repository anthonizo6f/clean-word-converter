import tempfile
"""生成舊式 .xls OLE 測試樣本（Real_XLS_OLE_Sample.docx）。

結構同真實 Excel.Sheet.8 嵌入一致：oleObject bin 本身就係成個 .xls CFB。
需要 LibreOffice（soffice）將 xlsx 轉 xls。
"""
import io
import os
import subprocess
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from make_fixture import OUT_DIR, build_cfb, make_xlsx_2

from docx import Document
from openpyxl import Workbook


def make_xls_bytes(workdir: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "BS"
    ws["A1"] = "XYZ Limited — 資產負債表"
    ws.merge_cells("A1:C1")
    ws.append([])
    ws.append(["項目", "2026 HK$", "2025 HK$"])
    ws.append(["固定資產", 987654.32, 876543.21])
    ws.append(["現金", 12345.67, ""])
    for addr in ("B4", "B5", "C4"):
        ws[addr].number_format = "#,##0.00"
    ws2 = wb.create_sheet("TC")
    ws2.append(["項目", "HK$"])
    ws2.append(["應評稅利潤", 250000])
    src = os.path.join(workdir, "_xls_src.xlsx")
    wb.save(src)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "xls", "--outdir", workdir, src],
        check=True, capture_output=True, timeout=120,
    )
    with open(src.replace(".xlsx", ".xls"), "rb") as f:
        return f.read()


def build_xls_fixture() -> str:
    os.makedirs(OUT_DIR, exist_ok=True)
    base = os.path.join(OUT_DIR, "_base_xls.docx")
    doc = Document()
    doc.add_paragraph("XLS_BEFORE 客戶提供嘅舊式 Excel 嵌入報表：")
    doc.add_paragraph("OLEOBJ1")
    doc.add_paragraph("XLS_AFTER 以上係舊格式，以下係新格式：")
    doc.add_paragraph("OLEOBJ2")
    doc.add_paragraph("XLS_END 完")
    doc.save(base)

    payloads = {1: make_xls_bytes(OUT_DIR), 2: build_cfb(make_xlsx_2())}
    progids = {1: "Excel.Sheet.8", 2: "Excel.Sheet.12"}
    out = os.path.join(OUT_DIR, "Real_XLS_OLE_Sample.docx")
    with zipfile.ZipFile(base) as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/document.xml":
                xml = data.decode("utf-8")
                for n in (1, 2):
                    old = f"<w:r><w:t>OLEOBJ{n}</w:t></w:r>"
                    new = ('<w:r><w:object xmlns:o="urn:schemas-microsoft-com:office:office">'
                           f'<o:OLEObject Type="Embed" ProgID="{progids[n]}" '
                           f'ShapeID="_x0000_i20{n}" DrawAspect="Content" '
                           f'ObjectID="_18000000{n}" r:id="rId20{n}"/></w:object></w:r>')
                    assert xml.count(old) == 1, f"marker {n} not found"
                    xml = xml.replace(old, new)
                data = xml.encode("utf-8")
            elif item == "word/_rels/document.xml.rels":
                xml = data.decode("utf-8")
                add = "".join(
                    f'<Relationship Id="rId20{n}" '
                    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" '
                    f'Target="embeddings/oleObject{n}.bin"/>' for n in (1, 2))
                data = xml.replace("</Relationships>", add + "</Relationships>").encode("utf-8")
            elif item == "[Content_Types].xml":
                xml = data.decode("utf-8")
                add = "".join(
                    f'<Override PartName="/word/embeddings/oleObject{n}.bin" '
                    'ContentType="application/vnd.openxmlformats-officedocument.oleObject"/>'
                    for n in (1, 2))
                data = xml.replace("</Types>", add + "</Types>").encode("utf-8")
            zout.writestr(item, data)
        for n, blob in payloads.items():
            zout.writestr(f"word/embeddings/oleObject{n}.bin", blob)
    os.remove(base)
    print("fixture:", out)
    return out


if __name__ == "__main__":
    build_xls_fixture()
