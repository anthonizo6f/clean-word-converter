"""
╔═══════════════════════════════════════════════════════════════╗
║  Clean Word Converter — 將 .docx 內嵌 Excel OLE 物件          ║
║  原位轉換為原生 Word 表格（IRD iXBRL 前置處理）                ║
║  Author: Anthonizo · Zo Computer                              ║
╚═══════════════════════════════════════════════════════════════╝

Usage (CLI):
    python clean_word_converter.py --input <file.docx | folder> [--output-dir DIR] [--log FILE]

Library:
    from clean_word_converter import convert_docx
    stats = convert_docx("in.docx", "out_clean.docx")
"""

from __future__ import annotations

import argparse
import datetime
import io
import logging
import os
import re
import shutil
import struct
import sys
import zipfile
from dataclasses import dataclass, field
from decimal import Decimal

import olefile
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string

log = logging.getLogger("clean_word_converter")

NS_O = "urn:schemas-microsoft-com:office:office"
RELTYPE_OLE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"

PLACEHOLDER_TEXT = "[TABLE EXTRACTION FAILED — manual review required]"
HARD_CAP_ROWS = 2000
HARD_CAP_COLS = 100
# ── 描述欄橫向合併（amendment 2：同事回饋）──────────────
# 將工作表「描述欄配對」（預設頭兩欄 A、B）視為同一描述區：
#   · 同一行只有 1 欄有內容 → 橫向合併兩欄，令描述跨欄顯示；
#   · 同一行第 1、第 2 欄都有內容 → 唔合併，保留兩格、格內自動換行。
# 已有 merge / overflow merge 覆蓋、或該格係數字格 → 一律唔郁。
DESCRIPTION_MERGE = True
DESCRIPTION_MERGE_PAIR = (0, 1)

# ══════════════════════════════════════════════════════════════
#  Data model
# ══════════════════════════════════════════════════════════════

@dataclass
class SheetData:
    name: str
    rows: list[list[str]] = field(default_factory=list)   # rendered cell text
    merges: list[tuple[int, int, int, int]] = field(default_factory=list)  # 0-based (r1,c1,r2,c2)
    col_widths: list[float] = field(default_factory=list)  # Excel 字元欄闊（可見欄）
    num_cells: set[tuple[int, int]] = field(default_factory=set)  # 數字/日期格（general 預設右對齊）
    cell_styles: dict[tuple[int, int], dict] = field(default_factory=dict)  # 邊框/字體/對齊
    row_heights: list[float | None] = field(default_factory=list)  # 原 Excel 行高（pt）
    gridlines: bool = True  # Excel「顯示格線」設定

    @property
    def n_rows(self) -> int:
        return len(self.rows)

    @property
    def n_cols(self) -> int:
        return max((len(r) for r in self.rows), default=0)


# ── Excel 視覺樣式擷取（邊框/字體/對齊）──────────────────────

_SBD = {"thin": ("single", "4"), "hair": ("single", "4"), "medium": ("single", "8"),
        "thick": ("single", "12"), "double": ("double", "4"),
        "dashed": ("dashed", "4"), "dotted": ("dotted", "4")}
_SBD_XLS = {0: None, 1: ("single", "4"), 2: ("single", "8"), 3: ("dashed", "4"),
            4: ("dotted", "4"), 5: ("single", "12"), 6: ("double", "4"), 7: ("single", "4"),
            8: ("dashed", "4"), 9: ("dotted", "4"), 10: ("dashed", "4"), 11: ("single", "8"),
            12: ("dotted", "4"), 13: ("dashed", "4")}
_ALIGN_XLS = {1: "left", 2: "center", 3: "right"}


def _capture_style(bold, italic, underline, edges: dict,
                   font_name=None, font_size=None, align=None) -> dict | None:
    """edges: {edge: (val, sz, color) | None}。全部預設 → 回傳 None。"""
    st: dict = {}
    if bold:
        st["bold"] = True
    if italic:
        st["italic"] = True
    if underline:
        st["underline"] = True
    bd = {e: v for e, v in edges.items() if v is not None}
    if bd:
        st["borders"] = bd
    if font_name:
        st["font_name"] = font_name
    if font_size:
        st["font_size"] = font_size
    if align in ("left", "center", "right"):
        st["align"] = align
    return st or None


def _xlsx_style(cell) -> dict | None:
    f = cell.font
    b = cell.border
    edges = {}
    for edge, side in (("top", b.top), ("bottom", b.bottom), ("left", b.left), ("right", b.right)):
        val = None
        if side is not None and side.style:
            color = "auto"
            try:
                if side.color and side.color.type == "rgb" and isinstance(side.color.rgb, str):
                    color = side.color.rgb[-6:]
            except Exception:
                pass
            val = (*_SBD.get(side.style, ("single", "4")), color)
        edges[edge] = val
    size = None
    try:
        if f.sz:
            size = float(f.sz)
    except Exception:
        pass
    return _capture_style(bool(f.b), bool(f.i), bool(f.u), edges,
                          font_name=f.name, font_size=size,
                          align=cell.alignment.horizontal)


def _xls_style(book, xf) -> dict | None:
    try:
        font = book.font_list[xf.font_index]
        bold = bool(font.bold)
        italic = bool(font.italic)
        underline = bool(getattr(font, "underline_type", 0))
        name = font.name or None
        size = round(font.height / 20.0, 1) if font.height else None
    except Exception:
        bold = italic = underline = False
        name = size = None
    edges = {}
    try:
        bd = xf.border
        for edge, attr in (("top", "top_line_style"), ("bottom", "bottom_line_style"),
                           ("left", "left_line_style"), ("right", "right_line_style")):
            m = _SBD_XLS.get(getattr(bd, attr, 0))
            edges[edge] = (*m, "auto") if m else None
    except Exception:
        edges = {"top": None, "bottom": None, "left": None, "right": None}
    align = None
    try:
        align = _ALIGN_XLS.get(xf.alignment.hor_align)
    except Exception:
        pass
    return _capture_style(bold, italic, underline, edges,
                          font_name=name, font_size=size, align=align)


class ExtractionError(Exception):
    """Raised when an OLE object cannot be converted to sheet data."""


# ══════════════════════════════════════════════════════════════
#  1. OLE binary → spreadsheet payload
# ══════════════════════════════════════════════════════════════

def _extract_ole_payload(ole_bytes: bytes) -> tuple[str, bytes]:
    """Return (kind, payload). kind ∈ {'xlsx', 'xls'}."""
    if not olefile.isOleFile(io.BytesIO(ole_bytes)):
        raise ExtractionError("嵌入物件不是有效的 OLE compound file")

    ole = olefile.OleFileIO(io.BytesIO(ole_bytes))
    try:
        streams = {"/".join(s) for s in ole.listdir(streams=True, storages=False)}

        raw = None
        for cand in ("\x01Ole10Native", "CONTENTS"):
            if ole.exists(cand):
                raw = ole.openstream(cand).read()
                break
        if raw is None:
            if any(s.endswith("Workbook") or s.endswith("/Book") for s in streams):
                return "xls", ole_bytes  # whole CFB is the xls (Excel.Sheet.8)
            raise ExtractionError(f"OLE 內找不到數據 stream（現有 streams: {sorted(streams)}）")

        # Ole10Native: 4-byte LE length prefix + native data
        if len(raw) >= 4:
            cb = struct.unpack("<I", raw[:4])[0]
            if 0 < cb <= len(raw) - 4:
                raw = raw[4:4 + cb]
            else:
                idx = raw.find(b"PK\x03\x04")
                if idx < 0:
                    idx = raw.find(b"\xd0\xcf\x11\xe0")
                raw = raw[idx:] if idx >= 0 else raw[4:]

        if raw[:4] == b"PK\x03\x04":
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                names = set(zf.namelist())
            if any(n.startswith("xl/") for n in names):
                return "xlsx", raw
            raise ExtractionError("嵌入的 ZIP 不是 Excel 活頁簿（可能是 Word/PDF 物件）")
        if raw[:4] == b"\xd0\xcf\x11\xe0":
            return "xls", raw
        raise ExtractionError("無法識別 OLE 內嵌數據格式（既非 xlsx 亦非 xls）")
    finally:
        ole.close()


# ══════════════════════════════════════════════════════════════
#  2. Payload → SheetData
# ══════════════════════════════════════════════════════════════

def _fmt_general_number(v) -> str:
    d = Decimal(str(v))
    s = format(d, "f")
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return s


def _split_format_sections(fmt: str) -> list[str]:
    """按 ; 分段（忽略引號內嘅 ;）。Excel 格式最多四段：正;負;零;文字。"""
    sections, buf, in_q = [], [], False
    for ch in fmt:
        if ch == '"':
            in_q = not in_q
        if ch == ";" and not in_q:
            sections.append("".join(buf))
            buf = []
        else:
            buf.append(ch)
    sections.append("".join(buf))
    return sections


def _literal_section_text(sec: str) -> str | None:
    """Section 無實數 placeholder（0/#）時回傳字面文字（如 "-"）；否則回傳 None。
    注意：? 只係對齊保留位（Accounting 格式常見 "-??"），唔當數字內容。"""
    s = re.sub(r"\[[^\]]*\]", "", sec)      # [Red] / [$-804] 等
    s = re.sub(r"[_*].", "", s)                # _) / * 填充 token
    no_q = re.sub(r'"[^"]*"', "", s)
    if re.search(r"[0#]", no_q):
        return None
    quoted = re.findall(r'"([^"]*)"', s)
    text = "".join(quoted) if quoted else "".join(ch for ch in no_q if ch in "-–—+() ").strip()
    if not text:
        return None
    # ? = Accounting 對齊預留位（每個佔一個數字位）→ 補 figure space（U+2007）
    return text + "\u2007" * no_q.count("?")


def _fmt_number(v: float, fmt: str) -> str:
    """Honour common Excel number formats（含 正;負;零 段式）；fall back to exact decimal."""
    try:
        sections = _split_format_sections(fmt)
        pick = sections[0]
        if v < 0 and len(sections) >= 2:
            pick = sections[1]
        elif v == 0 and len(sections) >= 3:
            pick = sections[2]
        lit = _literal_section_text(pick)
        if lit is not None:
            # amendment 1：數值 0 + Accounting 零值 dash（'-"??' / '"-"'）唔適合 tagging
            # → 一律輸出實數 0。其他字面 zero-section（如 "N/A"）保持原樣。
            if v == 0 and all(ch in "-–—\u2007" for ch in lit):
                return "0"
            return lit
        f = re.sub(r"\[[^\]]*\]", "", pick)
        f = re.sub(r'"[^"]*"', "", f)
        f = re.sub(r"[_*].", "", f)
        neg_paren = f.strip().startswith("(")
        f = f.replace("(", "").replace(")", "")
        av = abs(v)
        if "%" in f:
            dec = 0
            if "." in f:
                for ch in f.split(".")[1].split("%")[0]:
                    if ch in "0#?":
                        dec += 1
                    else:
                        break
            out = f"{av * 100:,.{dec}f}%"
        elif "." in f:
            dec = 0
            for ch in f.split(".")[1]:
                if ch in "0#?":
                    dec += 1
                else:
                    break
            out = f"{av:,.{dec}f}" if "," in f else f"{av:.{dec}f}"
        elif "," in f:
            out = f"{av:,.0f}"
        elif f.strip("0#? ") == "":
            out = f"{av:.0f}"
        else:
            return _fmt_general_number(v)
        if neg_paren:
            return f"({out})"
        if v < 0 and len(sections) == 1:
            return f"-{out}"
        return out
    except Exception:
        return _fmt_general_number(v)


def _render_cell(cell) -> str:
    v = cell.value
    if v is None:
        return ""
    if isinstance(v, bool):
        return "TRUE" if v else "FALSE"
    if isinstance(v, (datetime.datetime, datetime.date, datetime.time)):
        fmt = (cell.number_format or "").lower()
        if isinstance(v, datetime.time) or ("h" in fmt and "yyyy" not in fmt):
            return v.strftime("%H:%M:%S") if isinstance(v, datetime.time) else v.strftime("%Y-%m-%d %H:%M")
        return v.strftime("%Y-%m-%d %H:%M") if ("h" in fmt) else v.strftime("%Y-%m-%d")
    if isinstance(v, (int, float)):
        fmt = cell.number_format or "General"
        return _fmt_general_number(v) if fmt == "General" else _fmt_number(v, fmt)
    return str(v)


def _read_xlsx(payload: bytes, logger: logging.Logger) -> list[SheetData]:
    wb = load_workbook(io.BytesIO(payload), data_only=True, read_only=False)
    sheets: list[SheetData] = []
    for ws in wb.worksheets:
        max_r = min(ws.max_row or 0, HARD_CAP_ROWS)
        max_c = min(ws.max_column or 0, HARD_CAP_COLS)
        if (ws.max_row or 0) > HARD_CAP_ROWS or (ws.max_column or 0) > HARD_CAP_COLS:
            logger.warning("    ⚠ Sheet '%s' 超出上限，截斷至 %d 行 × %d 欄", ws.title, max_r, max_c)

        # ── 原 Excel 欄闊 + hidden 行列 + 行高 + 格線設定（還原視覺佈局）──
        col_hidden: set[int] = set()          # 1-based
        col_w_chars: dict[int, float] = {}
        for dim in ws.column_dimensions.values():
            try:
                lo = dim.min if dim.min is not None else column_index_from_string(str(dim.index))
                hi = dim.max if dim.max is not None else lo
            except Exception:
                continue
            for ci in range(lo, min(hi, HARD_CAP_COLS) + 1):
                if dim.hidden:
                    col_hidden.add(ci)
                if dim.width:
                    col_w_chars[ci] = float(dim.width)
        row_hidden = {int(ri) for ri, dim in ws.row_dimensions.items() if getattr(dim, "hidden", False)}
        row_ht = {int(ri): float(dim.height) for ri, dim in ws.row_dimensions.items()
                  if getattr(dim, "height", None)}
        gridlines = bool(ws.sheet_view.showGridLines)

        col_keep = [c for c in range(1, max_c + 1) if c not in col_hidden]   # 1-based
        col_map = {old: new for new, old in enumerate(col_keep)}             # 1-based → 0-based
        row_map: dict[int, int] = {}                                         # 1-based → 0-based

        grid: list[list[str]] = []
        num_cells: set[tuple[int, int]] = set()
        styles: dict[tuple[int, int], dict] = {}
        heights: list[float | None] = []
        last_r, last_c = -1, -1
        for r1 in range(1, max_r + 1):
            if r1 in row_hidden:
                continue
            new_r = len(grid)
            row_map[r1] = new_r
            heights.append(row_ht.get(r1))
            vals: list[str] = []
            for c1 in col_keep:
                cell = ws.cell(row=r1, column=c1)
                txt = _render_cell(cell)
                vals.append(txt)
                if isinstance(cell.value, (int, float, datetime.datetime, datetime.date)) \
                        and not isinstance(cell.value, bool):
                    num_cells.add((new_r, len(vals) - 1))
                st = _xlsx_style(cell)
                if st:
                    styles[(new_r, len(vals) - 1)] = st
                if txt != "":
                    last_r = new_r
                    last_c = max(last_c, len(vals) - 1)
            grid.append(vals)

        grid = [r[: last_c + 1] for r in grid[: last_r + 1]] if last_r >= 0 else []
        heights = heights[: last_r + 1] if last_r >= 0 else []
        num_cells = {(r, c) for (r, c) in num_cells if r <= last_r and c <= last_c}
        styles = {(r, c): s for (r, c), s in styles.items() if r <= last_r and c <= last_c}

        merges = []
        for mr in ws.merged_cells.ranges:
            rows_v = [r for r in range(mr.min_row, mr.max_row + 1) if r in row_map]
            cols_v = [c for c in range(mr.min_col, mr.max_col + 1) if c in col_map]
            if not rows_v or not cols_v:
                continue
            nr1, nr2 = row_map[rows_v[0]], row_map[rows_v[-1]]
            nc1, nc2 = col_map[cols_v[0]], col_map[cols_v[-1]]
            if nr1 <= last_r and nc1 <= last_c and (nr2 > nr1 or nc2 > nc1):
                merges.append((nr1, nc1, min(nr2, last_r), min(nc2, last_c)))

        default_w = ws.sheet_format.defaultColWidth or 8.43
        widths = [col_w_chars.get(c1, default_w) for c1 in col_keep[: last_c + 1]]

        sd = SheetData(name=ws.title, rows=grid, merges=merges,
                       col_widths=widths, num_cells=num_cells,
                       cell_styles=styles, row_heights=heights, gridlines=gridlines)
        df = pd.DataFrame(grid)
        logger.info("    · Sheet '%s': %s", sd.name, " × ".join(map(str, df.shape)) if grid else "空白")
        sheets.append(sd)
    wb.close()
    return sheets


def _read_xls(payload: bytes, logger: logging.Logger) -> list[SheetData]:
    try:
        import xlrd
    except ImportError as exc:
        raise ExtractionError("舊式 .xls 嵌入物件需要 xlrd（pip install xlrd）") from exc
    try:
        book = xlrd.open_workbook(file_contents=payload, formatting_info=True)
    except Exception:
        book = xlrd.open_workbook(file_contents=payload)  # 無格式資料都要照抽
    sheets = []
    for sh in book.sheets():
        colinfo = getattr(sh, "colinfo_map", {})
        rowinfo = getattr(sh, "rowinfo_map", {})
        col_keep = [c for c in range(sh.ncols) if not getattr(colinfo.get(c), "hidden", 0)]
        col_map = {old: new for new, old in enumerate(col_keep)}
        col_w = {c: colinfo[c].width / 256.0 for c in col_keep
                 if c in colinfo and getattr(colinfo[c], "width", 0)}
        gridlines = bool(getattr(sh, "show_grid_lines", True))

        def _cell(r, c, _sh=sh):
            ctype = _sh.cell_type(r, c)
            if ctype in (0, 6):  # empty / blank
                return ""
            v = _sh.cell_value(r, c)
            if ctype == 2:  # number — 套用原 Excel number format（含 正;負;零 段式）
                fmt = "General"
                try:
                    key = book.xf_list[_sh.cell_xf_index(r, c)].format_key
                    fmt = book.format_map[key].format_str
                except Exception:
                    pass
                return _fmt_general_number(v) if fmt == "General" else _fmt_number(v, fmt)
            if ctype == 3:  # date
                try:
                    dt = xlrd.xldate_as_datetime(v, book.datemode)
                    return dt.strftime("%Y-%m-%d") if dt.time() == datetime.time(0, 0) else dt.strftime("%Y-%m-%d %H:%M")
                except Exception:
                    return str(v)
            if ctype == 4:  # boolean
                return "TRUE" if v else "FALSE"
            return str(v)

        grid: list[list[str]] = []
        num_cells: set[tuple[int, int]] = set()
        styles: dict[tuple[int, int], dict] = {}
        heights: list[float | None] = []
        row_map: dict[int, int] = {}
        last_r, last_c = -1, -1
        for r in range(sh.nrows):
            rinfo = rowinfo.get(r)
            if rinfo is not None and getattr(rinfo, "hidden", 0):
                continue
            new_r = len(grid)
            row_map[r] = new_r
            ht = getattr(rinfo, "height", 0) if rinfo is not None else 0
            heights.append(round(ht / 20.0, 1) if ht else None)  # twips → pt
            vals: list[str] = []
            for c in col_keep:
                txt = _cell(r, c)
                vals.append(txt)
                if sh.cell_type(r, c) in (2, 3):
                    num_cells.add((new_r, len(vals) - 1))
                try:
                    xf = book.xf_list[sh.cell_xf_index(r, c)]
                    st = _xls_style(book, xf)
                    if st:
                        styles[(new_r, len(vals) - 1)] = st
                except Exception:
                    pass
                if txt != "":
                    last_r = new_r
                    last_c = max(last_c, len(vals) - 1)
            grid.append(vals)

        grid = [row[: last_c + 1] for row in grid[: last_r + 1]] if last_r >= 0 else []
        heights = heights[: last_r + 1] if last_r >= 0 else []
        num_cells = {(r, c) for (r, c) in num_cells if r <= last_r and c <= last_c}
        styles = {(r, c): s for (r, c), s in styles.items() if r <= last_r and c <= last_c}

        merges = []
        for (r1, r2, c1, c2) in getattr(sh, "merged_cells", []):  # xlrd: r2/c2 exclusive
            rows_v = [r for r in range(r1, r2) if r in row_map]
            cols_v = [c for c in range(c1, c2) if c in col_map]
            if not rows_v or not cols_v:
                continue
            nr1, nr2 = row_map[rows_v[0]], row_map[rows_v[-1]]
            nc1, nc2 = col_map[cols_v[0]], col_map[cols_v[-1]]
            if nr1 <= last_r and nc1 <= last_c and (nr2 > nr1 or nc2 > nc1):
                merges.append((nr1, nc1, min(nr2, last_r), min(nc2, last_c)))

        widths = [col_w.get(c, 8.43) for c in col_keep[: last_c + 1]]

        sheets.append(SheetData(name=sh.name, rows=grid, merges=merges,
                                col_widths=widths, num_cells=num_cells,
                                cell_styles=styles, row_heights=heights, gridlines=gridlines))
        logger.info("    · Sheet '%s' (xls): %d × %d", sh.name,
                    len(grid), len(grid[0]) if grid else 0)
    return sheets


# ══════════════════════════════════════════════════════════════
#  3. SheetData → native Word table elements
# ══════════════════════════════════════════════════════════════

MIN_COL_DXA = 450     # 欄闊下限（0.3"）
MAX_TABLE_DXA = 9000  # 後備頁寬；convert_docx 會按文件實際版心更新
_PAGE_CONTENT_DXA = MAX_TABLE_DXA


def _widths_to_dxa(widths_chars: list[float]) -> list[int]:
    """Excel 字元欄闊 → Word dxa，按比例放大/縮小至撐滿版心寬度。"""
    target = _PAGE_CONTENT_DXA
    raw = [max(int((w * 7 + 5) * 15), MIN_COL_DXA) for w in widths_chars]
    total = sum(raw)
    if total > 0 and total != target:
        scale = target / total
        raw = [max(int(w * scale), MIN_COL_DXA) for w in raw]
    return raw


def _text_width(s: str) -> int:
    """粗略顯示寬度：CJK/全形當 2，其他當 1（對應 Excel 欄寬單位）。"""
    return sum(2 if ord(ch) > 0x2E7F else 1 for ch in s)


def _overflow_merges(sd: SheetData) -> list[tuple[int, int, int, int]]:
    """Excel 顯示邏輯：文字格內容長過欄寬、而右邊係連續空格 → 溢出到隔籬格。
    Word 表格唔會自動溢出，所以將該格同右邊空格橫向合併嚟還原。
    規則：只處理文字格（數字唔溢出）；遇非空格/已有 merge 即停；每行獨立處理。"""
    if sd.n_rows == 0 or sd.n_cols == 0 or not sd.col_widths:
        return []
    merged = set()
    for (r1, c1, r2, c2) in sd.merges:
        for rr in range(r1, r2 + 1):
            for cc in range(c1, c2 + 1):
                merged.add((rr, cc))
    extra: list[tuple[int, int, int, int]] = []
    for r in range(sd.n_rows):
        row = sd.rows[r]
        c = 0
        while c < sd.n_cols:
            txt = row[c] if c < len(row) else ""
            if (not txt) or (r, c) in merged or (r, c) in sd.num_cells:
                c += 1
                continue
            if _text_width(txt) <= sd.col_widths[c]:
                c += 1
                continue
            # 文字太長 — 向右搵連續空格去合併
            end = c
            while end + 1 < sd.n_cols:
                nxt = row[end + 1] if end + 1 < len(row) else ""
                if nxt or (r, end + 1) in merged:
                    break
                end += 1
            if end > c:
                extra.append((r, c, r, end))
                for cc in range(c, end + 1):
                    merged.add((r, cc))
            c = end + 1
    return extra


def _set_table_chrome(tbl, widths_dxa: list[int] | None = None) -> None:
    """欄闊設定；底層框線永遠 nil — Excel OLE 預覽跟打印邏輯，唔顯示格線，
    只有真實 cell 邊框（由 _apply_cell_borders 逐格畫）。"""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)
    tblW = OxmlElement("w:tblW")
    if widths_dxa:
        tblW.set(qn("w:w"), str(sum(widths_dxa)))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
        tbl.autofit = False  # fixed layout — Word 唔會再自動拉均
        for gc, w in zip(tbl._tbl.tblGrid.findall(qn("w:gridCol")), widths_dxa):
            gc.set(qn("w:w"), str(w))
    else:
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)
        tbl.autofit = True


def _style_cell(cell) -> None:
    """Base 字體（無原 Excel 字體資料時嘅 fallback）。"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)


def _apply_cell_borders(cell, edge_map: dict) -> None:
    """逐邊寫入 w:tcBorders。edge_map: {edge: (val, sz, color)}。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge not in edge_map:
            continue
        val, sz, color = edge_map[edge]
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), sz)
            el.set(qn("w:color"), color)
        tb.append(el)
    tcPr.append(tb)


_ALIGN_WORD = {"left": WD_ALIGN_PARAGRAPH.LEFT,
               "center": WD_ALIGN_PARAGRAPH.CENTER,
               "right": WD_ALIGN_PARAGRAPH.RIGHT}


def _apply_cell_style(cell, st: dict) -> None:
    for para in cell.paragraphs:
        if st.get("align"):
            para.alignment = _ALIGN_WORD[st["align"]]
        for run in para.runs:
            if st.get("bold"):
                run.font.bold = True
            if st.get("italic"):
                run.font.italic = True
            if st.get("underline"):
                run.font.underline = True
            if st.get("font_name"):
                run.font.name = st["font_name"]
            if st.get("font_size"):
                run.font.size = Pt(st["font_size"])
    if "borders" in st:
        _apply_cell_borders(cell, st["borders"])


def _set_row_height(row, pt: float) -> None:
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(pt * 20)))  # pt → twips
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)



def _description_merges(sd: SheetData) -> list[tuple[int, int, int, int]]:
    """amendment 2：描述欄配對（預設頭兩欄 A、B）視為同一描述區。
    同一行只有 1 欄有內容 → 橫向合併兩欄；第 1、第 2 欄都有內容 → 唔合併（格內自動換行）。
    已有 sheet/overflow merge 覆蓋嘅格、或數字格 → 一律唔郁。"""
    if not DESCRIPTION_MERGE:
        return []
    c1, c2 = DESCRIPTION_MERGE_PAIR
    if sd.n_rows == 0 or c1 >= sd.n_cols or c2 >= sd.n_cols:
        return []
    covered = set()
    for (r1, cc1, r2, cc2) in sd.merges + _overflow_merges(sd):
        for rr in range(r1, r2 + 1):
            for cc in range(cc1, cc2 + 1):
                covered.add((rr, cc))
    extra: list[tuple[int, int, int, int]] = []
    for r in range(sd.n_rows):
        if (r, c1) in covered or (r, c2) in covered:
            continue
        if (r, c1) in sd.num_cells or (r, c2) in sd.num_cells:
            continue  # 數字格保持原樣，唔當描述合併
        row = sd.rows[r]
        t1 = row[c1] if c1 < len(row) else ""
        t2 = row[c2] if c2 < len(row) else ""
        has1, has2 = bool((t1 or "").strip()), bool((t2 or "").strip())
        if has1 == has2:  # 兩欄都有內容（各自自動換行）／兩欄都空 → 唔合併
            continue
        extra.append((r, c1, r, c2))
    return extra


def _build_sheet_elements(doc: Document, sd: SheetData) -> list:
    """Return [table_element]；空白 sheet 直接略過（唔加任何痕跡）。"""
    if sd.n_rows == 0 or sd.n_cols == 0:
        return []

    widths_dxa = (_widths_to_dxa(sd.col_widths)
                  if sd.col_widths and len(sd.col_widths) == sd.n_cols else None)
    tbl = doc.add_table(rows=sd.n_rows, cols=sd.n_cols)
    _set_table_chrome(tbl, widths_dxa)
    _c1, _c2 = DESCRIPTION_MERGE_PAIR
    _right_only: dict[int, str] = {}
    if DESCRIPTION_MERGE and _c1 < sd.n_cols and _c2 < sd.n_cols:
        _covered = set()
        for (_r1, _cc1, _r2, _cc2) in sd.merges + _overflow_merges(sd):
            for _rr in range(_r1, _r2 + 1):
                for _cc in range(_cc1, _cc2 + 1):
                    _covered.add((_rr, _cc))
        for _r, _row in enumerate(sd.rows):
            if (_r, _c1) in _covered or (_r, _c2) in _covered:
                continue
            if (_r, _c1) in sd.num_cells or (_r, _c2) in sd.num_cells:
                continue
            _t2 = _row[_c2] if _c2 < len(_row) else ""
            if not (_row[_c1] if _c1 < len(_row) else "").strip() and (_t2 or "").strip():
                _right_only[_r] = _t2
    for r, row in enumerate(sd.rows):
        if r < len(sd.row_heights) and sd.row_heights[r]:
            _set_row_height(tbl.rows[r], sd.row_heights[r])
        for c in range(sd.n_cols):
            txt = row[c] if c < len(row) else ""
            if DESCRIPTION_MERGE and c == _c1 and r in _right_only:
                txt = _right_only[r]
            elif DESCRIPTION_MERGE and c == _c2 and r in _right_only:
                txt = ""
            cell = tbl.cell(r, c)
            cell.text = txt
            _style_cell(cell)
            if widths_dxa:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(widths_dxa[c]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)
            if (r, c) in sd.num_cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            st = sd.cell_styles.get((r, c))
            if st:
                _apply_cell_style(cell, st)  # 原 Excel 設定優先（會覆蓋上面嘅預設對齊）
    _desc_merges = _description_merges(sd)
    for (r1, c1, r2, c2) in sd.merges + _overflow_merges(sd) + _desc_merges:
        if r2 < sd.n_rows and c2 < sd.n_cols and (r2 > r1 or c2 > c1):
            try:
                merged = tbl.cell(r1, c1).merge(tbl.cell(r2, c2))
                # merge 會把被併入儲存格的空段落一併保留 → 清走，只留首個有文字的段落
                kept = False
                for pel in merged._tc.findall(qn("w:p")):
                    has_text = any((t.text or "").strip()
                                   for t in pel.findall(".//" + qn("w:t")))
                    if has_text and not kept:
                        kept = True
                        continue
                    if has_text or kept:
                        merged._tc.remove(pel)
            except Exception:
                pass  # overlapping merge — keep unmerged, values retained
    return [tbl._tbl]


def _placeholder_paragraph(doc: Document, reason: str) -> "element":
    p = doc.add_paragraph()
    run = p.add_run(PLACEHOLDER_TEXT)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    if reason:
        r2 = p.add_run(f"  （{reason}）")
        r2.font.name = "Calibri"
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p._p


# ══════════════════════════════════════════════════════════════
#  4. Document surgery — copy & splice
# ══════════════════════════════════════════════════════════════

def _ancestor_run(el):
    node = el.getparent()
    while node is not None and node.tag != qn("w:r"):
        node = node.getparent()
    return node


def _paragraph_is_empty(p) -> bool:
    for t in p.findall(".//" + qn("w:t")):
        if t.text and t.text.strip():
            return False
    for tag in ("w:drawing", "w:pict", "w:object"):
        if p.findall(".//" + qn(tag)):
            return False
    return True


def convert_docx(input_path: str, output_path: str, logger: logging.Logger | None = None) -> dict:
    logger = logger or log
    stats = {
        "file": os.path.basename(input_path),
        "output_path": output_path,
        "objects_found": 0,
        "objects_replaced": 0,
        "objects_failed": 0,
        "tables_created": 0,
        "ole_remaining": -1,
        "size_mb": 0.0,
        "status": "ok",
        # ── Report 用 fields ──
        "paragraphs_before": 0, "paragraphs_after": 0,
        "tables_before": 0, "tables_after": 0,
        "text_chars_before": 0, "text_chars_after": 0,
        "tables_metadata": [],
        "integrity": "ok",
        "warnings": [],
    }
    logger.info("處理檔案：%s", input_path)

    # ── 輸入文件統計 ──
    try:
        before_doc = Document(input_path)
        body = before_doc.element.body
        stats["paragraphs_before"] = len(body.findall(".//" + qn("w:p")))
        stats["tables_before"] = len(body.findall(".//" + qn("w:tbl")))
        stats["text_chars_before"] = sum(len(t.text or "") for t in body.findall(".//" + qn("w:t")))
    except Exception:
        pass

    shutil.copyfile(input_path, output_path)
    doc = Document(output_path)

    global _PAGE_CONTENT_DXA
    try:
        sec = doc.sections[0]
        _PAGE_CONTENT_DXA = int((sec.page_width - sec.left_margin - sec.right_margin) / 635)
    except Exception:
        _PAGE_CONTENT_DXA = MAX_TABLE_DXA

    ole_rels = {rid: rel for rid, rel in doc.part.rels.items() if rel.reltype == RELTYPE_OLE}
    stats["objects_found"] = len(ole_rels)
    logger.info("  找到 %d 個 OLE 嵌入物件（版心寬 %d dxa）", len(ole_rels), _PAGE_CONTENT_DXA)

    if not ole_rels:
        doc.save(output_path)
        stats["ole_remaining"] = 0
        stats["size_mb"] = round(os.path.getsize(output_path) / 1e6, 3)
        logger.info("  無 OLE 物件，原樣複製輸出")
        return stats

    processed_rids: set[str] = set()

    for p in doc.element.body.findall(".//" + qn("w:p")):
        ole_objs = p.findall(".//{%s}OLEObject" % NS_O)
        if not ole_objs:
            continue
        anchor = p
        for ole in ole_objs:
            rid = ole.get(qn("r:id"))
            processed_rids.add(rid)
            try:
                rel = ole_rels[rid]
                payload_kind, payload = _extract_ole_payload(rel.target_part.blob)
                sheets = (_read_xlsx if payload_kind == "xlsx" else _read_xls)(payload, logger)
                for sd in sheets:
                    stats.setdefault("tables_metadata", []).append({
                        "sheet_name": sd.name,
                        "rows": sd.n_rows,
                        "cols": sd.n_cols,
                        "header": sd.rows[0] if sd.n_rows > 0 else [],
                        "ole_object": rid,
                    })
                    for el in _build_sheet_elements(doc, sd):
                        anchor.addnext(el)
                        anchor = el
                        if el.tag == qn("w:tbl"):
                            stats["tables_created"] += 1
                stats["objects_replaced"] += 1
                logger.info("  ✓ OLE 物件 %s → %d 個工作表已轉換", rid, len(sheets))
            except Exception as exc:  # noqa: BLE001 — per-object isolation
                reason = str(exc)[:120]
                logger.error("  ✗ OLE 物件 %s 抽取失敗：%s", rid, reason)
                el = _placeholder_paragraph(doc, reason)
                anchor.addnext(el)
                anchor = el
                stats["objects_failed"] += 1
            run = _ancestor_run(ole)
            if run is not None:
                run.getparent().remove(run)
        if _paragraph_is_empty(p):
            p.getparent().remove(p)

    unprocessed = set(ole_rels) - processed_rids
    if unprocessed:
        logger.warning("  ⚠ %d 個 OLE 物件位於表格/頁首頁尾等未支援位置，未轉換：%s",
                       len(unprocessed), sorted(unprocessed))

    for rid in processed_rids:
        try:
            doc.part.drop_rel(rid)
        except Exception:
            pass

    doc.save(output_path)

    with zipfile.ZipFile(output_path) as zf:
        remaining_parts = [n for n in zf.namelist() if n.startswith("word/embeddings/")]
        xml = zf.read("word/document.xml").decode("utf-8", "ignore")
    stats["ole_remaining"] = xml.count("OLEObject") + len(remaining_parts)
    stats["size_mb"] = round(os.path.getsize(output_path) / 1e6, 3)
    if stats["ole_remaining"]:
        logger.warning("  ⚠ 輸出仍有 %d 項 OLE 殘留", stats["ole_remaining"])
    # ── 輸出文件統計（跟輸入比對）──
    try:
        after_doc = Document(output_path)
        abody = after_doc.element.body
        stats["paragraphs_after"] = len(abody.findall(".//" + qn("w:p")))
        stats["tables_after"] = len(abody.findall(".//" + qn("w:tbl")))
        stats["text_chars_after"] = sum(len(t.text or "") for t in abody.findall(".//" + qn("w:t")))
    except Exception:
        pass

    # ── integrity / warnings ──
    # 注意：段落/文字數喺轉換後必然增加（OLE 二元碼 → 原生表格文字），
    # 真正嘅 integrity 指標係 OLE 處理成功與否，唔係前後數量差異。
    w = stats.setdefault("warnings", [])
    if stats["objects_failed"]:
        w.append(f"{stats['objects_failed']} OLE extraction failed")
    if stats["ole_remaining"] > 0:
        w.append(f"{stats['ole_remaining']} OLE residuals remaining")
    if stats["text_chars_after"] < stats["text_chars_before"]:
        w.append("text content loss detected (after < before)")
    stats["integrity"] = ("fail" if stats["objects_found"] > 0 and stats["objects_replaced"] == 0
                          else "partial" if w else "ok")

    logger.info("  完成：%d 成功 / %d 失敗，輸出 %s（%.2f MB）integrity=%s",
                stats["objects_replaced"], stats["objects_failed"], output_path, stats["size_mb"],
                stats["integrity"])
    if stats["objects_failed"]:
        stats["status"] = "partial"
    return stats


# ══════════════════════════════════════════════════════════════
#  5. CLI
# ══════════════════════════════════════════════════════════════

def _setup_logger(log_path: str) -> logging.Logger:
    logger = logging.getLogger("clean_word_converter")
    logger.setLevel(logging.INFO)
    logger.handlers.clear()
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s", "%Y-%m-%d %H:%M:%S")
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setFormatter(fmt)
    sh = logging.StreamHandler(sys.stdout)
    sh.setFormatter(fmt)
    logger.addHandler(fh)
    logger.addHandler(sh)
    return logger


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="將 .docx 內嵌 Excel OLE 物件轉換為原生 Word 表格（IRD iXBRL 前置處理）")
    ap.add_argument("--input", required=True, help="輸入 .docx 檔案或資料夾")
    ap.add_argument("--output-dir", default="./output", help="輸出資料夾（預設 ./output）")
    ap.add_argument("--log", default=None, help="Log 檔路徑（預設 <output-dir>/clean_word_converter.log）")
    ap.add_argument("--report", action="store_true",
                    help="輸出轉換驗證報告（_report.json + _report.csv）")
    args = ap.parse_args(argv)

    os.makedirs(args.output_dir, exist_ok=True)
    log_path = args.log or os.path.join(args.output_dir, "clean_word_converter.log")
    logger = _setup_logger(log_path)

    if os.path.isdir(args.input):
        targets = sorted(
            os.path.join(args.input, f) for f in os.listdir(args.input)
            if f.lower().endswith(".docx") and not f.startswith("~$")
        )
    else:
        targets = [args.input]

    if not targets:
        logger.error("找不到任何 .docx 檔案：%s", args.input)
        return 1

    logger.info("═══ Clean Word Converter 開始 · %d 個檔案 ═══", len(targets))
    results = []
    for path in targets:
        out_path = os.path.join(
            args.output_dir, os.path.splitext(os.path.basename(path))[0] + "_clean.docx")
        try:
            results.append(convert_docx(path, out_path, logger))
        except Exception as exc:  # noqa: BLE001 — per-file isolation
            logger.error("檔案層面失敗 %s：%s", path, exc)
            results.append({"file": os.path.basename(path), "status": "error",
                            "objects_found": 0, "objects_replaced": 0,
                            "objects_failed": 0, "tables_created": 0,
                            "ole_remaining": -1, "size_mb": 0, "output_path": ""})

    # ── 輸出驗證報告（--report）──
    if args.report:
        import json, csv
        report_json = os.path.join(args.output_dir, "conversion_report.json")
        report_csv  = os.path.join(args.output_dir, "conversion_report.csv")
        for r in results:
            r.setdefault("integrity", "error" if r["status"] == "error" else "ok")
            r.setdefault("warnings", [])
        with open(report_json, "w", encoding="utf-8") as fj:
            json.dump(results, fj, ensure_ascii=False, indent=2, default=str)
        logger.info("驗證報告 JSON 已輸出：%s", report_json)
        # CSV — flat 視圖（排除 tables_metadata 巢狀欄）
        csv_fields = ["file", "output_path", "objects_found", "objects_replaced",
                      "objects_failed", "tables_created", "ole_remaining", "size_mb",
                      "paragraphs_before", "paragraphs_after", "tables_before",
                      "tables_after", "text_chars_before", "text_chars_after",
                      "integrity", "status", "warnings"]
        with open(report_csv, "w", encoding="utf-8-sig", newline="") as fc:
            cw = csv.DictWriter(fc, fieldnames=csv_fields, extrasaction="ignore")
            cw.writeheader()
            for r in results:
                row = {k: r.get(k, "") for k in csv_fields}
                if isinstance(row.get("warnings"), list):
                    row["warnings"] = " | ".join(str(v) for v in row["warnings"])
                cw.writerow(row)
        logger.info("驗證報告 CSV  已輸出：%s", report_csv)

    logger.info("═══ 完成 ═══")
    for r in results:
        logger.info("  %s → 找到 %d OLE · 成功 %d · 失敗 %d · 表格 %d · %.2f MB [%s]",
                    r["file"], r["objects_found"], r["objects_replaced"],
                    r["objects_failed"], r["tables_created"], r["size_mb"], r["status"])
    return 1 if any(r["status"] == "error" for r in results) else 0


if __name__ == "__main__":
    raise SystemExit(main())
