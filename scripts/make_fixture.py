"""Build a realistic .docx test fixture with genuine embedded Excel OLE objects."""
import io
import os
import struct
import sys
import tempfile
import zipfile

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

OUT_DIR = sys.argv[1] if len(sys.argv) > 1 else os.path.join(tempfile.gettempdir(), "cwc_fixtures")
os.makedirs(OUT_DIR, exist_ok=True)

FREE, END, FATSECT = 0xFFFFFFFF, 0xFFFFFFFE, 0xFFFFFFFD


def build_cfb(native: bytes) -> bytes:
    ole10 = struct.pack("<I", len(native)) + native
    size = max(4096, -(-len(ole10) // 512) * 512)
    ole10 = ole10.ljust(size, b"\x00")
    ndata = size // 512

    fat = [FREE] * 128
    fat[0], fat[1] = FATSECT, END
    for i in range(ndata):
        fat[2 + i] = 2 + i + 1 if i < ndata - 1 else END

    def direntry(name, typ, start, sz, child=FREE):
        enc = name.encode("utf-16-le")
        e = (enc + b"\x00\x00").ljust(64, b"\x00")
        e += struct.pack("<H", len(enc) + 2)
        e += struct.pack("<BB", typ, 1)
        e += struct.pack("<III", FREE, FREE, child)
        e += b"\x00" * 16 + struct.pack("<I", 0) + b"\x00" * 16
        e += struct.pack("<I", start) + struct.pack("<Q", sz)
        return e

    free_e = (b"\x00" * 64 + struct.pack("<HBB", 0, 0, 1)
              + struct.pack("<III", FREE, FREE, FREE) + b"\x00" * 16
              + struct.pack("<I", 0) + b"\x00" * 16
              + struct.pack("<I", FREE) + struct.pack("<Q", 0))
    dir_bytes = (direntry("Root Entry", 5, END, 0, child=1)
                 + direntry("\x01Ole10Native", 2, 2, size) + free_e + free_e)

    hdr = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 16
    hdr += struct.pack("<HHHHH", 0x003E, 3, 0xFFFE, 9, 6) + b"\x00" * 6
    hdr += struct.pack("<IIII", 0, 1, 1, 0)
    hdr += struct.pack("<I", 4096)
    hdr += struct.pack("<IIII", END, 0, END, 0)
    hdr += struct.pack("<109I", 0, *([FREE] * 108))
    return hdr + struct.pack("<128I", *fat) + dir_bytes + ole10


def make_xlsx_1() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Balance Sheet"
    ws["A1"] = "ABC Company Limited — 資產負債表 Balance Sheet"
    ws.merge_cells("A1:D1")
    ws.append([])
    ws.append(["項目 Item", "附註 Note", "2026 HK$", "2025 HK$"])
    ws.append(["非流動資產 Non-current assets", "", None, None])
    r5 = ["物業、機器及設備 PP&E", "4", 1234567.89, 1100000]
    ws.append(r5)
    ws["C5"].number_format = "#,##0.00"
    ws["D5"].number_format = "#,##0.00"
    ws.append(["流動資產 Current assets", "", 654321, None])
    ws["C6"].number_format = "#,##0"
    ws.append(["純利率 Net margin", "", 0.156, 0.132])
    ws["C7"].number_format = "0.0%"
    ws["D7"].number_format = "0.0%"
    ws.append(["合計 Total", "", 1888888.89, None])
    ws["C8"].number_format = "#,##0.00"
    ws.merge_cells("A8:B8")
    # — 視覺還原測試：零值三段式、負數括弧、自訂欄闊、hidden column —
    ws.append(["董事袍金 Director's fees", "", 0, 0])
    ws["C9"].number_format = '#,##0;(#,##0);"-"'
    ws["D9"].number_format = '#,##0_);(#,##0);"-"_)'
    ws.append(["調整項 Adjustment", "", -1234.5, None])
    ws["C10"].number_format = "#,##0.00;(#,##0.00)"
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 6
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 16
    ws["E3"] = "HIDDEN_COL"
    ws["E5"] = 999
    ws.column_dimensions["E"].hidden = True
    # — 視覺樣式測試：header 底線、total 行頂線+粗體、自訂行高、serif 字體、右對齊 header、長 label —
    from openpyxl.styles import Border, Side
    thin_bottom = Border(bottom=Side(style="thin"))
    for col in ("A", "B", "C", "D"):
        ws[f"{col}3"].border = thin_bottom
        ws[f"{col}3"].font = Font(name="Times New Roman", size=12, bold=True)
        if col in ("C", "D"):
            ws[f"{col}3"].alignment = Alignment(horizontal="right")
    ws.append(["INCOME", "", None, None])
    ws["A11"].font = Font(name="Times New Roman", size=12, bold=True)
    ws.append(["COST OF SALES", "", None, None])
    ws["A12"].font = Font(name="Times New Roman", size=12, bold=True)
    top_border = Border(top=Side(style="thin"))
    for col in ("A", "B", "C", "D"):
        ws[f"{col}8"].border = top_border
        ws[f"{col}8"].font = Font(bold=True)
    ws.row_dimensions[1].height = 30
    ws.sheet_view.showGridLines = True
    ws3 = wb.create_sheet("Overflow")
    ws3.column_dimensions["A"].width = 6
    ws3.column_dimensions["B"].width = 8
    ws3.column_dimensions["C"].width = 10
    ws3.append(["INCOME", "", None])
    ws3["A1"].font = Font(bold=True)
    ws3.append(["COST OF SALES", None, None])
    ws3.append(["Note", "5", 1000])
    ws3["C3"].number_format = "#,##0"
    ws3.sheet_view.showGridLines = False
    ws2 = wb.create_sheet("Notes")
    ws2.sheet_view.showGridLines = False
    ws2.append(["Note", "內容 Content"])
    ws2.append(["4", "物業、機器及設備按成本減累計折舊列賬"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_xlsx_2() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Tax Computation"
    ws.append(["項目 Item", "HK$"])
    ws.append(["Assessable profits 應評稅利潤", 500000])
    ws["B2"].number_format = "#,##0"
    ws.append(["Tax payable @16.5%", 82500])
    ws["B3"].number_format = "#,##0.00"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_docx():
    base = os.path.join(OUT_DIR, "_base.docx")
    doc = Document()
    doc.add_paragraph("TEXT_BEFORE 本公司截至2026年3月31日止年度之財務報表。")
    doc.add_paragraph("OLEOBJ1")
    doc.add_paragraph("TEXT_BETWEEN 上述報表已按香港財務報告準則編製。")
    doc.add_paragraph("OLEOBJ2")
    doc.add_paragraph("TEXT_BAD_LEAD 下表因檔案損毀未能自動抽取：")
    doc.add_paragraph("OLEOBJ3")
    doc.add_paragraph("TEXT_AFTER 謹啟")
    doc.save(base)

    payloads = {
        1: build_cfb(make_xlsx_1()),
        2: build_cfb(make_xlsx_2()),
        3: b"\x00\x01GARBAGE-NOT-AN-OLE-FILE",
    }
    out = os.path.join(OUT_DIR, "Embedded_OLE_Sample.docx")
    with zipfile.ZipFile(base) as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/document.xml":
                xml = data.decode("utf-8")
                for n in (1, 2, 3):
                    old = f"<w:r><w:t>OLEOBJ{n}</w:t></w:r>"
                    new = (
                        '<w:r><w:object xmlns:o="urn:schemas-microsoft-com:office:office">'
                        f'<o:OLEObject Type="Embed" ProgID="Excel.Sheet.12" '
                        f'ShapeID="_x0000_i10{n}" DrawAspect="Content" '
                        f'ObjectID="_17000000{n}" r:id="rId10{n}"/></w:object></w:r>'
                    )
                    assert xml.count(old) == 1, f"marker {n} not found"
                    xml = xml.replace(old, new)
                data = xml.encode("utf-8")
            elif item == "word/_rels/document.xml.rels":
                xml = data.decode("utf-8")
                add = "".join(
                    f'<Relationship Id="rId10{n}" '
                    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" '
                    f'Target="embeddings/oleObject{n}.bin"/>' for n in (1, 2, 3))
                data = xml.replace("</Relationships>", add + "</Relationships>").encode("utf-8")
            elif item == "[Content_Types].xml":
                xml = data.decode("utf-8")
                add = "".join(
                    f'<Override PartName="/word/embeddings/oleObject{n}.bin" '
                    'ContentType="application/vnd.openxmlformats-officedocument.oleObject"/>'
                    for n in (1, 2, 3))
                data = xml.replace("</Types>", add + "</Types>").encode("utf-8")
            zout.writestr(item, data)
        for n, blob in payloads.items():
            zout.writestr(f"word/embeddings/oleObject{n}.bin", blob)
    os.remove(base)
    print("fixture:", out)
    return out


if __name__ == "__main__":
    build_docx()
