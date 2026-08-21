"""Clean Word Converter 最終視覺回歸測試（自包含）。

用法：
    python final_visual_test.py

會自動生成 fixture → 轉換 → 斷言全部視覺還原原點：
欄寬、零值 dash、Accounting dash 補位、負數括弧、間線（底層 nil + 真實邊框）、
粗體/字體名/磅數/對齊、行高、hidden 欄、overflow 橫向合併、無 Sheet: 標題段、零 OLE 殘留。
"""
import os
import re
import subprocess
import sys
import zipfile

HERE = os.path.dirname(os.path.abspath(__file__))
SKILL = os.path.dirname(HERE)
sys.path.insert(0, HERE)

FIX = os.path.join(SKILL, "assets", "_fixtures_build")
OUT = os.path.join(FIX, "_out")

import make_fixture  # noqa: E402
import build_xls_fixture  # noqa: E402
import clean_word_converter as m  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

# ── 生成 fixture ──
make_fixture.OUT_DIR = FIX
os.makedirs(FIX, exist_ok=True)
make_fixture.build_docx()
build_xls_fixture.OUT_DIR = FIX
build_xls_fixture.build_xls_fixture()


def edge(cell, e):
    pr = cell._tc.find(qn("w:tcPr"))
    if pr is None:
        return None
    tb = pr.find(qn("w:tcBorders"))
    if tb is None:
        return None
    el = tb.find(qn(f"w:{e}"))
    return (el.get(qn("w:val")), el.get(qn("w:sz"))) if el is not None else None


# ── 轉換 ──
os.makedirs(OUT, exist_ok=True)
for fn in ("Embedded_OLE_Sample.docx", "Real_XLS_OLE_Sample.docx"):
    m.convert_docx(os.path.join(FIX, fn),
                   os.path.join(OUT, fn.replace(".docx", "_clean.docx")))

# ═══ xlsx fixture ═══
doc = Document(f"{OUT}/Embedded_OLE_Sample_clean.docx")
assert len(doc.tables) == 4, f"應 4 表（BS/Overflow/Notes/TaxComp），實得 {len(doc.tables)}"
t0, t1, t2, t3 = doc.tables

assert t0.cell(8, 2).text == "0" and t0.cell(8, 3).text == "0"  # amendment 1：零值顯示實數 0
assert t0.cell(9, 2).text == "(1,234.50)"
assert t0.cell(4, 2).text == "1,234,567.89" and t0.cell(4, 3).text == "1,100,000.00"
assert t0.cell(6, 2).text == "15.6%" and t0.cell(6, 3).text == "13.2%"
assert len(t0.columns) == 4
assert all("HIDDEN" not in t0.cell(i, c).text for i in range(len(t0.rows)) for c in range(4))
print("✅ 數值格式（零值 - / 負數括弧 / 千分位 / % / hidden 欄）")

acc = m._fmt_number(0, '_(* #,##0.00_);_(* (#,##0.00);_(* "-"??_)')
assert acc == "0", f"零值應顯示實數 0（amendment 1），實得 {acc!r}"
assert m._fmt_number(0, '#,##0;(#,##0);"-"') == "0"
print("✅ 零值 Accounting dash → 實數 0（適合 tagging）")# 負數唔可以變正數（同事 amendment：舊版 $ 前綴括號負數段被忽略）
assert m._fmt_number(-1234.5, '_($* #,##0.00_);_($* (#,##0.00);_($* "-"??_);_(@_)') == "(1,234.50)"
assert m._fmt_number(-1234.5, '#,##0.00;(#,##0.00)') == "(1,234.50)"
assert m._fmt_number(-1234.5, '#,##0.00;[Red]-#,##0.00') == "-1,234.50"
print("✅ 負數括號段（$ 前綴）／顯式負號段：唔會再變正數")

grid = [int(g.get(qn("w:w"))) for g in t0._tbl.tblGrid.findall(qn("w:gridCol"))]
assert grid[0] > grid[2] > grid[1] and abs(grid[2] - grid[3]) <= 25 and sum(grid) > 7000
print(f"✅ 欄寬比例 {grid}（總寬 {sum(grid)} dxa）")

for t, nm in ((t0, "BS"), (t1, "Overflow"), (t2, "Notes"), (t3, "TaxComp")):
    tb = t._tbl.tblPr.find(qn("w:tblBorders"))
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        assert tb.find(qn(f"w:{e}")).get(qn("w:val")) == "nil", f"{nm} 底層 {e} 應 nil"
print("✅ 全部表格底層框線 = nil（跟 OLE 預覽邏輯）")
assert edge(t0.cell(2, 2), "bottom") == ("single", "4")
assert edge(t0.cell(7, 2), "top") == ("single", "4")
assert edge(t0.cell(7, 2), "bottom") is None
print("✅ 真實 cell 邊框（header 底線 / total 頂線），無多餘間線")

hdr = t0.cell(2, 2).paragraphs[0].runs[0]
assert hdr.font.bold and hdr.font.name == "Times New Roman" and hdr.font.size.pt == 12.0
assert t0.cell(7, 2).paragraphs[0].runs[0].font.bold
inc = t0.cell(10, 0).paragraphs[0].runs[0]
assert inc.font.name == "Times New Roman" and inc.font.size.pt == 12.0
print("✅ 粗體 + 字體名 + 磅數跟原 Excel")

assert t0.cell(4, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
assert t0.cell(2, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
assert t0.cell(4, 0).paragraphs[0].alignment is None
print("✅ 對齊（數字右 / header 右跟 Excel / label 左）")

tr0 = t0.rows[0]._tr.find(qn("w:trPr"))
assert tr0.find(qn("w:trHeight")).get(qn("w:val")) == "600"
print("✅ 行高跟原 Excel")

spans = [int(el.get(qn("w:val"))) for el in t1._tbl.findall(".//" + qn("w:gridSpan"))]
assert 3 in spans, f"長 label 應橫跨：{spans}"
assert t1.cell(1, 0).text == "COST OF SALES"
assert t1.cell(2, 0).text == "Note" and t1.cell(2, 2).text == "1,000"
print(f"✅ overflow 橫向合併（橫跨 {max(spans)} 欄，數據行完整）")

assert "Sheet:" not in "\n".join(p.text for p in doc.paragraphs)
print("✅ 無 Sheet: 標題段")

# ═══ xls fixture ═══
doc2 = Document(f"{OUT}/Real_XLS_OLE_Sample_clean.docx")
x0 = doc2.tables[0]
assert x0.cell(0, 0).text == "XYZ Limited — 資產負債表"
assert x0.cell(3, 1).text == "987,654.32" and x0.cell(4, 1).text == "12,345.67"
assert x0.cell(4, 2).text == ""
assert doc2.tables[1].cell(1, 1).text == "250000"
tb = x0._tbl.tblPr.find(qn("w:tblBorders"))
assert tb.find(qn("w:top")).get(qn("w:val")) == "nil"
print("✅ xls fixture（數值/格式/合併/無底層格線）")

for fn in ("Embedded_OLE_Sample_clean.docx", "Real_XLS_OLE_Sample_clean.docx"):
    with zipfile.ZipFile(os.path.join(OUT, fn)) as zf:
        assert "OLEObject" not in zf.read("word/document.xml").decode()
        assert not any("embeddings/oleObject" in n for n in zf.namelist())
print("✅ 零 OLE 殘留")

# ═══ Report 功能單元測試 ═══
REPORT_DIR = os.path.join(FIX, "_report_test")
os.makedirs(REPORT_DIR, exist_ok=True)
import subprocess as _sp  # noqa: E402
_sp.run([sys.executable, os.path.join(HERE, "clean_word_converter.py"),
         "--input", FIX, "--output-dir", REPORT_DIR, "--report"],
        check=True, capture_output=True)
import json, csv  # noqa: E402
with open(os.path.join(REPORT_DIR, "conversion_report.json"), encoding="utf-8") as _fj:
    _reports = json.load(_fj)
_xls_r = [r for r in _reports if "XLS" in r["file"]][0]
_xlsx_r = [r for r in _reports if "Embedded_OLE" in r["file"]][0]
assert _xls_r["integrity"] == "ok", f"xls integrity={_xls_r['integrity']}"
assert _xlsx_r["integrity"] == "partial", f"xlsx integrity={_xlsx_r['integrity']}"
assert _xls_r["objects_failed"] == 0
assert _xlsx_r["objects_failed"] == 1
assert len(_xls_r["tables_metadata"]) == 3
assert len(_xlsx_r["tables_metadata"]) == 4
_tm = _xlsx_r["tables_metadata"]
assert any(m["sheet_name"] == "Overflow" and m["rows"] == 3 for m in _tm)
assert any(m["sheet_name"] == "Balance Sheet" for m in _tm)
assert os.path.exists(os.path.join(REPORT_DIR, "conversion_report.csv"))
print("✅ Report：integrity ok/partial 正確 + tables_metadata 計到 4 張表")

# ── amendment 2：描述欄橫向合併 ──
def _gs(cell):
    pr = cell._tc.find(qn("w:tcPr"))
    if pr is None:
        return None
    el = pr.find(qn("w:gridSpan"))
    return int(el.get(qn("w:val"))) if el is not None else None

# 單欄有內容 → 跨兩欄合併；兩欄都有內容 → 各自換行；不與既有/overflow merge 衝突
descs = {3: "非流動資產 Non-current assets", 5: "流動資產 Current assets",
         8: "董事袍金 Director's fees", 10: "INCOME", 11: "COST OF SALES"}
for _r, _label in descs.items():
    assert _gs(t0.cell(_r, 0)) == 2, f"row {_r} 單欄描述應跨兩欄合併"
    assert t0.cell(_r, 0).text == _label, f"row {_r} 文字應保留：{_label!r}"
assert _gs(t0.cell(4, 0)) is None and t0.cell(4, 1).text == "4"  # 項目+附註都有 → 唔合併
assert _gs(t0.cell(7, 0)) == 2                                   # 原本已 merge 嘅 合計 唔重覆
assert _gs(t1.cell(1, 0)) == 3                                   # overflow 合併優先，唔受影響
print("✅ 描述欄 merge：單欄跨兩欄 / 兩欄各自換行 / 不與 overflow 衝突")

# ── amendment 2 單元測試 ──
_d1 = Document()
m._build_sheet_elements(_d1, m.SheetData(name="T", rows=[["", "右欄獨有描述"]], col_widths=[20.0, 8.0]))
_t = _d1.tables[0]
assert _t.cell(0, 0).text == "右欄獨有描述" and _gs(_t.cell(0, 0)) == 2   # right-only → 搬左 + 合併
_d2 = Document()
m._build_sheet_elements(_d2, m.SheetData(name="T", rows=[["左欄", "右欄"]], col_widths=[20.0, 8.0]))
_t2 = _d2.tables[0]
assert _t2.cell(0, 0).text == "左欄" and _t2.cell(0, 1).text == "右欄" and _gs(_t2.cell(0, 0)) is None
_d3 = Document()
m._build_sheet_elements(_d3, m.SheetData(name="T", rows=[["", "123"]], col_widths=[20.0, 8.0], num_cells={(0, 1)}))
_t3 = _d3.tables[0]
assert _t3.cell(0, 1).text == "123" and _gs(_t3.cell(0, 0)) is None   # 數字格 → 唔合併
print("✅ right-only 搬左合併 / 兩欄都有各自換行 / 數字格唔郁")
print("\n🎯 FINAL VISUAL ASSERTIONS ALL PASSED")
